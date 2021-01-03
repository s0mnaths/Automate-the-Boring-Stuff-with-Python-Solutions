import docx

doc=docx.Document()
file=open('guests.txt')
NamesList=file.readlines()
for Name in NamesList:
    doc.add_paragraph('It would be a pleasure to have the company of', style='Custom 1')
    doc.add_paragraph(Name.strip(), style='Custom 2')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='Custom 3')
    doc.add_paragraph('April 1st', style='Custom 4')
    doc.add_paragraph("at 7 o'clock", style='Custom 5')
    doc.add_page_break()

doc.save('textToInvitations.docx')
print('Done!')